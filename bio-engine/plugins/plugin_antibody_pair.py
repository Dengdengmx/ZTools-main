# plugins/plugin_antibody_pair.py
"""
抗体序列智能配对与模板化输出插件
支持从多个 Word (DOCX) 或 Excel (XLSX) 文档中自动提取重链(H)/轻链(L)，
依据核心编号自动配对，并在前端预览后，依照指定 Excel 模板的样式生成最终报告。
"""

import os
import re
import sys
import traceback
from collections import Counter, defaultdict
from copy import copy
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple

from docx import Document
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

from PyQt5.QtWidgets import (QWidget, QVBoxLayout, QHBoxLayout, QFileDialog, 
                             QMessageBox, QListWidgetItem, QAbstractItemView, QHeaderView, QSplitter)
from PyQt5.QtCore import Qt

from qfluentwidgets import (PrimaryPushButton, PushButton, TextEdit, BodyLabel, 
                            StrongBodyLabel, ListWidget, LineEdit, TableWidget, FluentIcon as FIF)

from core.plugin_manager import BasePlugin, PluginWorker
from core.ui_base import BasePluginUI
from core.config import GlobalConfig


# ==========================================
# 底层数据结构与正则引擎
# ==========================================
@dataclass
class SeqRecord:
    raw_header: str
    short_name: str
    core_name: str
    chain_type: str   # H or L
    seq: str
    unique_name: str = ""

HEADER_PREFIX_RE = re.compile(r"^>")
# ✨ 修复 1：放宽正则限制，支持 "_123H_"，也支持最普通的 "123H"
CORE_NAME_RE = re.compile(r"(?:^|_)([A-Za-z0-9\-]+?)([HL])(?:_|$)")
DNA_RE = re.compile(r"^[ACGTUNacgtun\s]+$")

def clean_seq(seq: str) -> str:
    seq = re.sub(r"\s+", "", seq).upper()
    seq = seq.replace("U", "T")
    return seq

def parse_header(header: str) -> Tuple[str, str, str]:
    header = header.strip()
    if header.startswith(">"): header = header[1:].strip()
    m = CORE_NAME_RE.search(header)
    if not m: raise ValueError(f"无法从标题中识别 H/L 链特征：{header}")
    core_plus_nochain = m.group(1)
    chain_type = m.group(2)
    short_name = f"{core_plus_nochain}{chain_type}"
    core_name = core_plus_nochain
    return short_name, core_name, chain_type

def iter_text_blocks_from_docx(docx_path: str) -> Iterable[str]:
    doc = Document(docx_path)
    for p in doc.paragraphs:
        text = p.text.strip()
        if text: yield text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text: yield text

# ✨ 修复 2：赋予 Excel 智能双列识别能力，自动补全 FASTA 格式的 ">"
def iter_text_blocks_from_xlsx(xlsx_path: str) -> Iterable[str]:
    wb = load_workbook(xlsx_path, data_only=True)
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        for row in ws.iter_rows(values_only=True):
            # 如果这一行前两列都有数据，大概率是 [名称, 序列] 的排版
            if len(row) >= 2 and row[0] and row[1] and isinstance(row[0], str) and isinstance(row[1], str):
                name, seq = row[0].strip(), row[1].strip()
                # 校验第二列是否是纯序列
                if len(seq) > 5 and re.match(r"^[ACGTUNacgtun\s]+$", seq):
                    if not name.startswith(">"): name = f">{name}" # 智能补全
                    yield name
                    yield seq
                    continue
            
            # 兜底：逐个单元格提取
            for cell in row:
                if cell is not None:
                    text = str(cell).strip()
                    if text: yield text


# ==========================================
# 前端：UI 交互视图
# ==========================================
class AntibodyPairUI(BasePluginUI):
    def __init__(self, parent=None, is_setting_mode=False):
        super().__init__(plugin_id="antibody_pair", plugin_name="序列配对与模板输出", parent=parent)
        self.is_setting_mode = is_setting_mode
        self.file_list = []
        self.paired_results = []
        self.unpaired_h = []
        self.unpaired_l = []
        self.heavy_dups = []
        self.light_dups = []
        self.setAcceptDrops(True)
        self._setup_ui()
        self._load_config()

    def _setup_ui(self):
        # ========== 左侧参数面板 ==========
        if not self.is_setting_mode:
            group_data, layout_data = self.create_group("1. 输入文档队列 (支持 DOCX / XLSX 拖拽)") # ✨ 更新文案
            layout_data.setSpacing(4)
            btn_row = QHBoxLayout(); btn_row.setContentsMargins(0, 0, 0, 0)
            btn_add = PrimaryPushButton("导入序列文件...", icon=FIF.DOWNLOAD) # ✨ 更新文案
            btn_add.clicked.connect(self.open_input_files)
            btn_clear = PushButton("清空", icon=FIF.DELETE)
            btn_clear.clicked.connect(self.clear_files)
            btn_row.addWidget(btn_add); btn_row.addWidget(btn_clear)
            layout_data.addLayout(btn_row)

            self.list_widget = ListWidget()
            self.list_widget.setFixedHeight(120)
            self.list_widget.setSelectionMode(QAbstractItemView.ExtendedSelection)
            layout_data.addWidget(self.list_widget)
            
            btn_remove = PushButton("移除选中项", icon=FIF.REMOVE)
            btn_remove.clicked.connect(self.remove_selected_file)
            layout_data.addWidget(btn_remove)
            self.add_param_widget(group_data)

        group_tmpl, layout_tmpl = self.create_group("2. Excel 样式模板设置")
        layout_tmpl.setSpacing(4)
        
        self.line_template = LineEdit()
        self.line_template.setPlaceholderText("留空则生成默认排版 Excel")
        btn_tmpl = PushButton("浏览...")
        btn_tmpl.clicked.connect(self.select_template)
        self.add_row(layout_tmpl, "模板:", self.line_template, "", btn_tmpl)
        self.add_param_widget(group_tmpl)

        if not self.is_setting_mode:
            group_export, layout_export = self.create_group("3. 数据输出")
            layout_export.setSpacing(4)
            
            self.line_out_dir = LineEdit()
            btn_out_dir = PushButton("更改...")
            btn_out_dir.clicked.connect(self.select_out_dir)
            self.add_row(layout_export, "导出至:", self.line_out_dir, "", btn_out_dir)
            
            self.line_basename = LineEdit()
            self.line_basename.setText("Antibody_Pairing_Result")
            self.add_row(layout_export, "主文件名:", self.line_basename)
            self.add_param_widget(group_export)

        self.add_param_stretch()

        if self.is_setting_mode:
            self.btn_save_config = PrimaryPushButton("💾 保存配置", icon=FIF.SAVE)
            self.btn_save_config.clicked.connect(self.save_settings_and_close)
            self.param_layout.addWidget(self.btn_save_config)
            return

        # 核心操作按钮
        action_layout = QHBoxLayout()
        self.btn_parse = PushButton("🔍 解析并预览", icon=FIF.SEARCH)
        self.btn_parse.clicked.connect(self.run_parse)
        self.btn_export = PrimaryPushButton("💾 导出 Excel", icon=FIF.SAVE)
        self.btn_export.clicked.connect(self.run_export)
        self.btn_export.setEnabled(False)
        action_layout.addWidget(self.btn_parse)
        action_layout.addWidget(self.btn_export)
        self.param_layout.addLayout(action_layout)

        # ========== 右侧预览与日志画板 ==========
        splitter = QSplitter(Qt.Vertical)
        splitter.setStyleSheet("QSplitter::handle { background-color: #EAEAEA; height: 1px; }")

        # 1. 结果预览表格
        table_container = QWidget()
        table_layout = QVBoxLayout(table_container)
        table_layout.setContentsMargins(0, 0, 0, 0)
        table_layout.addWidget(StrongBodyLabel("👁️ 核心配对预览 (Core Pairing Preview)"))
        
        self.table = TableWidget(self)
        self.table.setColumnCount(4)
        self.table.setHorizontalHeaderLabels(['重链名称 (Heavy)', '重链序列 (部分)', '轻链名称 (Light)', '轻链序列 (部分)'])
        self.table.verticalHeader().hide()
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        table_layout.addWidget(self.table)
        splitter.addWidget(table_container)

        # 2. 日志控制台
        log_container = QWidget()
        log_layout = QVBoxLayout(log_container)
        log_layout.setContentsMargins(0, 5, 0, 0)
        log_layout.addWidget(StrongBodyLabel("📝 解析日志 (Logs & Warnings)"))
        
        self.log_text = TextEdit()
        self.log_text.setReadOnly(True)
        self.log_text.setStyleSheet("QTextEdit { background-color: #2D2D2D; color: #E0E0E0; font-family: Consolas; font-size: 11px; border-radius: 5px; padding: 5px; }")
        log_layout.addWidget(self.log_text)
        splitter.addWidget(log_container)
        
        splitter.setSizes([450, 150])
        self.get_canvas_layout().addWidget(splitter)
        self.log("系统就绪。请拖入需要解析的 Word 或 Excel 文档。")

    # --- 交互与事件 ---
    def log(self, msg, color="#E0E0E0"):
        self.log_text.append(f"<span style='color:{color};'>{msg}</span>")
        
    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls(): event.accept()
        else: event.ignore()

    def dropEvent(self, event):
        for url in event.mimeData().urls():
            fp = url.toLocalFile()
            # ✨ 允许拖拽 .xlsx 文件
            if fp.lower().endswith(('.docx', '.xlsx')) and fp not in self.file_list:
                self.file_list.append(fp)
                self.list_widget.addItem(os.path.basename(fp))

    def open_input_files(self):
        # ✨ 开放弹窗对 .xlsx 的支持
        files, _ = QFileDialog.getOpenFileNames(self, "选择序列文档", "", "Supported Files (*.docx *.xlsx);;Word Files (*.docx);;Excel Files (*.xlsx)")
        for fp in files:
            if fp not in self.file_list:
                self.file_list.append(fp)
                self.list_widget.addItem(os.path.basename(fp))

    def remove_selected_file(self):
        for item in self.list_widget.selectedItems():
            idx = self.list_widget.row(item)
            self.list_widget.takeItem(idx)
            self.file_list.pop(idx)

    def clear_files(self):
        self.file_list.clear()
        self.list_widget.clear()
        self.table.setRowCount(0)
        self.btn_export.setEnabled(False)

    def select_template(self):
        fp, _ = QFileDialog.getOpenFileName(self, "选择 Excel 模板", "", "Excel Files (*.xlsx)")
        if fp: self.line_template.setText(fp)

    def select_out_dir(self):
        folder = QFileDialog.getExistingDirectory(self, "选择输出目录")
        if folder: self.line_out_dir.setText(folder)

    # --- 配置管理 ---
    def _load_config(self):
        config = GlobalConfig.get_all_plugin_settings(self.plugin_id)
        if config:
            self.line_template.setText(config.get("template", ""))
            self.line_out_dir.setText(config.get("out_dir", os.getcwd()))

    def _save_config(self):
        GlobalConfig.set_all_plugin_settings(self.plugin_id, {
            "template": self.line_template.text(),
            "out_dir": self.line_out_dir.text()
        })

    def save_settings_and_close(self):
        self._save_config()
        self.window().accept()

    # ==========================================
    # 核心业务逻辑
    # ==========================================
    def run_parse(self):
        if not self.file_list:
            QMessageBox.warning(self, "提示", "队列中没有文件。")
            return
            
        self._save_config()
        self.log_text.clear()
        self.table.setRowCount(0)
        self.log(f"开始扫描 {len(self.file_list)} 个文档...")
        self.btn_parse.setEnabled(False)

        worker = PluginWorker(self._do_parse_logic, self.file_list)
        worker.signals.finished.connect(self._on_parse_finished)
        worker.signals.error.connect(self._on_parse_error)
        worker.signals.log.connect(self.log)
        
        from PyQt5.QtCore import QThreadPool
        QThreadPool.globalInstance().start(worker)

    def _do_parse_logic(self, file_paths, progress_callback=None, log_callback=None):
        all_records = []
        
        # 1. 全局混读解析
        for fp in file_paths:
            ext = os.path.splitext(fp)[1].lower()
            # ✨ 根据后缀名智能切换文本提取引擎
            if ext == '.docx':
                texts = list(iter_text_blocks_from_docx(fp))
            elif ext == '.xlsx':
                texts = list(iter_text_blocks_from_xlsx(fp))
            else:
                continue

            current_header, current_seq_parts = None, []

            def flush_seq():
                if not current_header: return
                seq = clean_seq("".join(current_seq_parts))
                if not seq:
                    log_callback.emit(f"⚠️ 跳过: {current_header} (无序列数据)")
                    return
                if not DNA_RE.match(seq):
                    bad = re.sub(r"[ACGTUNacgtun]", "", seq)
                    log_callback.emit(f"⚠️ 警告: 序列包含非标准字符 ({current_header}) -> {bad[:10]}")
                try:
                    short_name, core_name, chain_type = parse_header(current_header)
                    all_records.append(SeqRecord(current_header, short_name, core_name, chain_type, seq))
                except Exception as e:
                    log_callback.emit(f"⚠️ 命名不规范，已跳过: {current_header}")

            for text in texts:
                lines = [x.strip() for x in text.splitlines() if x.strip()]
                for line in lines:
                    if HEADER_PREFIX_RE.match(line):
                        flush_seq()
                        current_header = line
                        current_seq_parts = []
                    else:
                        if current_header: current_seq_parts.append(line)
            flush_seq()

        if not all_records:
            raise ValueError("所有文档均未提取到符合规范的序列数据！")

        # 2. 分配唯一名
        counts = Counter(r.short_name for r in all_records)
        running = Counter()
        for r in all_records:
            if counts[r.short_name] == 1: r.unique_name = r.short_name
            else:
                running[r.short_name] += 1
                r.unique_name = f"{r.short_name}-{running[r.short_name]}"

        heavy_records = [r for r in all_records if r.chain_type == 'H']
        light_records = [r for r in all_records if r.chain_type == 'L']
        
        # ✅ 修复 3：通过内嵌 HTML 样式实现变色，完美适配单参数信号槽
        log_callback.emit(f"<span style='color:#14CC14;'>✅ 成功提取：{len(heavy_records)} 条重链，{len(light_records)} 条轻链。</span>")

        # 3. 核心配对逻辑
        heavy_map = self._group_by_core(heavy_records)
        light_map = self._group_by_core(light_records)
        all_cores = sorted(set(heavy_map) | set(light_map))

        paired, up_h, up_l = [], [], []
        for core in all_cores:
            hs, ls = heavy_map.get(core, []), light_map.get(core, [])
            if hs and ls:
                for h in hs:
                    for l in ls:
                        paired.append({"heavy_name": h.unique_name, "heavy_seq": h.seq, "light_name": l.unique_name, "light_seq": l.seq})
            elif hs:
                for h in hs: up_h.append({"名称": h.unique_name, "核心编号": h.core_name, "序列": h.seq})
            elif ls:
                for l in ls: up_l.append({"名称": l.unique_name, "核心编号": l.core_name, "序列": l.seq})

        heavy_dups = self._find_duplicates(heavy_records)
        light_dups = self._find_duplicates(light_records)

        return (paired, up_h, up_l, heavy_dups, light_dups)

    def _group_by_core(self, records):
        d = defaultdict(list)
        for r in records: d[r.core_name].append(r)
        return d

    def _find_duplicates(self, records):
        seq_to_rec = defaultdict(list)
        for r in records: seq_to_rec[r.seq].append(r)
        dups, gid = [], 1
        for seq, rs in seq_to_rec.items():
            if len(rs) > 1:
                dups.append({
                    "重复组编号": gid, "出现次数": len(rs),
                    "名称列表": "; ".join([x.unique_name for x in rs]), "序列": seq
                })
                gid += 1
        return dups

    def _on_parse_finished(self, result):
        self.btn_parse.setEnabled(True)
        self.paired_results, self.unpaired_h, self.unpaired_l, self.heavy_dups, self.light_dups = result
        
        self.table.setRowCount(len(self.paired_results))
        from PyQt5.QtWidgets import QTableWidgetItem
        for i, row in enumerate(self.paired_results):
            self.table.setItem(i, 0, QTableWidgetItem(row["heavy_name"]))
            self.table.setItem(i, 1, QTableWidgetItem(row["heavy_seq"][:15] + "..."))
            self.table.setItem(i, 2, QTableWidgetItem(row["light_name"]))
            self.table.setItem(i, 3, QTableWidgetItem(row["light_seq"][:15] + "..."))
            
        self.log(f"🔗 配对完成: 成功组装 {len(self.paired_results)} 对序列。")
        if self.unpaired_h or self.unpaired_l:
            self.log(f"⚠️ 孤立序列: 未配对 H 链 ({len(self.unpaired_h)})，未配对 L 链 ({len(self.unpaired_l)})", "#FF5722")
            
        self.btn_export.setEnabled(True)

    def _on_parse_error(self, err_tuple):
        self.btn_parse.setEnabled(True)
        self.log(f"❌ 解析崩溃: {err_tuple[1]}", "#FF0000")

    # ==========================================
    # 模板化导出引擎
    # ==========================================
    def run_export(self):
        out_dir = self.line_out_dir.text()
        if not out_dir or not os.path.exists(out_dir):
            QMessageBox.warning(self, "错误", "请指定正确的输出目录。")
            return
            
        basename = self.line_basename.text() or "Pairing_Result"
        out_main = os.path.join(out_dir, f"{basename}.xlsx")
        out_unpaired = os.path.join(out_dir, f"{basename}_Unpaired.xlsx")
        
        tmpl = self.line_template.text()
        template_path = tmpl if os.path.exists(tmpl) else None

        try:
            self._write_main_workbook(out_main, template_path)
            if self.unpaired_h or self.unpaired_l:
                self._write_unpaired_workbook(out_unpaired)
                
            self.log(f"✅ 导出成功: {out_main}", "#14CC14")
            
            import platform, subprocess
            if platform.system() == "Windows": subprocess.Popen(f'explorer /select,"{os.path.normpath(out_main)}"')
            elif platform.system() == "Darwin": subprocess.Popen(['open', '-R', out_main])
                
        except Exception as e:
            self.log(f"❌ 导出失败: {str(e)}", "#FF0000")
            traceback.print_exc()

    def _write_main_workbook(self, out_xlsx, template_path):
        if template_path:
            template_wb = load_workbook(template_path)
            wb = Workbook()
            wb.remove(wb.active)
            for sheet_name in template_wb.sheetnames:
                src_ws = template_wb[sheet_name]
                dst_ws = wb.create_sheet(title=sheet_name)
                self._copy_template_sheet(src_ws, dst_ws)
        else:
            wb = Workbook()
            ws = wb.active
            ws.title = "序列-核酸or氨基酸"

        if "序列-核酸or氨基酸" in wb.sheetnames:
            ws = wb["序列-核酸or氨基酸"]
            max_row = ws.max_row
            if max_row >= 3:
                for row in ws.iter_rows(min_row=3, max_row=max_row, min_col=1, max_col=max(4, ws.max_column)):
                    for cell in row: cell.value = None
        else:
            ws = wb.active
            ws.title = "序列-核酸or氨基酸"

        if ws["A1"].value is None and ws["C1"].value is None:
            ws.merge_cells("A1:B1"); ws.merge_cells("C1:D1")
            ws["A1"] = "heavy"; ws["C1"] = "Light"
            ws["A2"] = "名称"; ws["B2"] = "序列"; ws["C2"] = "名称"; ws["D2"] = "序列"
            
            fill_blue = PatternFill("solid", fgColor="D9EAF7")
            fill_dark = PatternFill("solid", fgColor="B7DEE8")
            for cell in ws[1]: cell.font = Font(bold=True); cell.fill = fill_dark; cell.alignment = Alignment(horizontal="center", vertical="center")
            for cell in ws[2]: cell.font = Font(bold=True); cell.fill = fill_blue; cell.alignment = Alignment(horizontal="center", vertical="center")
            ws.column_dimensions["A"].width = 24
            ws.column_dimensions["B"].width = 42
            ws.column_dimensions["C"].width = 24
            ws.column_dimensions["D"].width = 42
            ws.freeze_panes = "A3"

        row_idx = 3
        for row in self.paired_results:
            ws.cell(row_idx, 1, row["heavy_name"])
            ws.cell(row_idx, 2, row["heavy_seq"])
            ws.cell(row_idx, 3, row["light_name"])
            ws.cell(row_idx, 4, row["light_seq"])
            row_idx += 1

        for r in range(3, row_idx):
            for c in range(1, 5): ws.cell(r, c).alignment = Alignment(vertical="top", wrap_text=True)

        self._write_dup_sheet(wb, "重链重复序列", self.heavy_dups)
        self._write_dup_sheet(wb, "轻链重复序列", self.light_dups)
        wb.save(out_xlsx)

    def _copy_template_sheet(self, src_ws, dst_ws):
        for row in src_ws.iter_rows():
            for cell in row:
                new_cell = dst_ws[cell.coordinate]
                new_cell.value = cell.value
                if cell.has_style: new_cell._style = copy(cell._style)
                if cell.font: new_cell.font = copy(cell.font)
                if cell.fill: new_cell.fill = copy(cell.fill)
                if cell.border: new_cell.border = copy(cell.border)
                if cell.alignment: new_cell.alignment = copy(cell.alignment)
                if cell.number_format: new_cell.number_format = copy(cell.number_format)
        for idx, dim in src_ws.row_dimensions.items(): dst_ws.row_dimensions[idx].height = dim.height
        for key, dim in src_ws.column_dimensions.items(): dst_ws.column_dimensions[key].width = dim.width; dst_ws.column_dimensions[key].hidden = dim.hidden
        for merged in src_ws.merged_cells.ranges: dst_ws.merge_cells(str(merged))

    def _write_dup_sheet(self, wb, title, rows):
        if title in wb.sheetnames: del wb[title]
        dws = wb.create_sheet(title=title)
        headers = ["重复组编号", "出现次数", "名称列表", "序列"]
        dws.append(headers)
        for x in rows: dws.append([x[h] for h in headers])
        for cell in dws[1]: cell.font = Font(bold=True); cell.fill = PatternFill("solid", fgColor="D9EAF7"); cell.alignment = Alignment(horizontal="center", vertical="center")
        self._autosize_columns(dws)

    def _write_unpaired_workbook(self, out_xlsx):
        wb = Workbook()
        ws1 = wb.active; ws1.title = "未配对重链"
        headers = ["名称", "核心编号", "序列"]
        ws1.append(headers)
        for x in self.unpaired_h: ws1.append([x[h] for h in headers])

        ws2 = wb.create_sheet("未配对轻链")
        ws2.append(headers)
        for x in self.unpaired_l: ws2.append([x[h] for h in headers])

        for ws in [ws1, ws2]:
            for cell in ws[1]: cell.font = Font(bold=True); cell.fill = PatternFill("solid", fgColor="D9EAF7"); cell.alignment = Alignment(horizontal="center", vertical="center")
            self._autosize_columns(ws)
        wb.save(out_xlsx)

    def _autosize_columns(self, ws, min_width=15, max_width=60):
        for col_cells in ws.columns:
            col_idx = col_cells[0].column
            letter = get_column_letter(col_idx)
            max_len = max([len(str(c.value)) if c.value else 0 for c in col_cells])
            ws.column_dimensions[letter].width = min(max(max_len * 1.5, min_width), max_width)

# ==========================================
# 插件底层注册
# ==========================================
class AntibodyPairPlugin(BasePlugin):
    plugin_id = "antibody_pair"
    plugin_name = "序列配对与模板化导出"
    icon = "🔗"
    trigger_tag = "抗体配对"

    def get_ui(self, parent=None):
        return AntibodyPairUI(parent, is_setting_mode=False)

    @staticmethod
    def run(file_path, archive_dir):
        return "", "【后端引擎提示】此插件目前主要依靠工作台交互生成批量矩阵，不直接参与自动化出图流水线。"